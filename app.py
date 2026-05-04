"""
Medical Device Requirements Extractor - Local Web App
Rallis-Daw Consulting | RegCheck Tool
------------------------------------------------------
Supports PDF, Word (.docx), and Excel (.xlsx) files.
Extracts regulatory standards using Gemini AI and stores
results in Supabase.

SETUP:
    pip3 install flask google-generativeai pdfplumber python-docx openpyxl python-dotenv supabase

CREATE A .env FILE in this same folder containing:
    GEMINI_API_KEY=your_gemini_api_key
    SUPABASE_URL=your_supabase_project_url
    SUPABASE_KEY=your_supabase_anon_key

RUN:
    python3 app.py

Then open your browser to: http://localhost:5000
"""

import json
import os
import re
import tempfile
from datetime import datetime, timezone
import pdfplumber
import google.generativeai as genai
from flask import Flask, request, jsonify, render_template_string
from docx import Document
from dotenv import load_dotenv
from supabase import create_client, Client
import openpyxl

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
SUPABASE_URL   = os.getenv("SUPABASE_URL")
SUPABASE_KEY   = os.getenv("SUPABASE_KEY")
GEMINI_MODEL   = "gemini-2.5-flash"
TABLE_NAME     = "extracted_from_List"

app = Flask(__name__)

# Initialize Supabase client
supabase_client = None
if SUPABASE_URL and SUPABASE_KEY:
    try:
        supabase_client = create_client(SUPABASE_URL, SUPABASE_KEY)
    except Exception as e:
        print(f"  WARNING: Supabase connection failed: {e}")


# ── TEXT EXTRACTION ────────────────────────────────────────────────────────────

def extract_from_pdf(path):
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_from_docx(path):
    doc = Document(path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += row_text + "\n"
    return text


def extract_from_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    text = ""
    for sheet in wb.worksheets:
        text += "\n--- Sheet: " + sheet.title + " ---\n"
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                text += row_text + "\n"
    return text


def extract_text(path, filename):
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_from_pdf(path)
    elif ext == "docx":
        return extract_from_docx(path)
    elif ext == "xlsx":
        return extract_from_xlsx(path)
    else:
        raise ValueError("Unsupported file type: ." + ext)


# ── DATA HELPERS ───────────────────────────────────────────────────────────────

def extract_year(date_str):
    if not date_str or date_str.strip().lower() in ("current", "**", ""):
        return None
    match = re.search(r'\b(19|20)\d{2}\b', str(date_str))
    return int(match.group()) if match else None


def enrich_requirement(req, filename):
    date_val = req.get("date", "") or ""
    is_current = date_val.strip().lower() in ("current", "**", "")
    year = extract_year(date_val)
    return {
        "standard_id":         req.get("standard_id", ""),
        "date":                date_val,
        "date_year":           int(year) if year is not None else None,
        "category":            req.get("category", ""),
        "region":              req.get("region", ""),
        "description":         req.get("description", ""),
        "needs_manual_review": is_current,
        "source_filename":     filename,
        "uploaded_at":         datetime.now(timezone.utc).isoformat(),
        "status":              None,
    }


def insert_to_supabase(requirements):
    if not supabase_client:
        return 0, "Supabase is not configured."
    try:
        supabase_client.table(TABLE_NAME).insert(requirements).execute()
        return len(requirements), None
    except Exception as e:
        return 0, str(e)


# ── HTML ───────────────────────────────────────────────────────────────────────

HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>RegCheck &mdash; Rallis-Daw Consulting</title>
<link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@400;500;600;700;800;900&display=swap" rel="stylesheet">
<style>
  :root {
    --navy:    #1a2f5e;
    --navy2:   #223570;
    --blue:    #2e6db4;
    --lblue:   #5ba3d9;
    --cyan:    #4dc8e8;
    --bg:      #eaf3fb;
    --white:   #ffffff;
    --text:    #1a2f5e;
    --muted:   #6b82a8;
    --border:  #c5d8ed;
    --success: #2d7a4f;
    --danger:  #c0392b;
    --warning: #b7770d;
    --font:    'Montserrat', sans-serif;
  }

  * { box-sizing: border-box; margin: 0; padding: 0; }

  body {
    background: var(--bg);
    color: var(--text);
    font-family: var(--font);
    min-height: 100vh;
    overflow-x: hidden;
  }

  .page {
    min-height: 100vh;
    position: relative;
    display: flex;
    flex-direction: column;
  }

  .wave-bg {
    position: fixed;
    top: 0; right: 0;
    width: 55%;
    height: 100vh;
    pointer-events: none;
    z-index: 0;
  }

  header {
    position: relative;
    z-index: 10;
    background: var(--white);
    padding: 16px 48px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    border-bottom: 1px solid var(--border);
    box-shadow: 0 2px 12px rgba(26,47,94,0.08);
  }

  .logo-wrap { display: flex; align-items: center; gap: 12px; }
  .logo-img { height: 52px; width: auto; }

  .tool-label {
    font-size: 11px;
    font-weight: 700;
    color: var(--muted);
    letter-spacing: 0.14em;
    text-transform: uppercase;
    padding: 3px 10px;
    border: 1px solid var(--border);
    border-radius: 3px;
  }

  .content {
    position: relative;
    z-index: 5;
    flex: 1;
    display: flex;
    flex-direction: column;
    padding: 48px 56px 60px;
    max-width: 780px;
  }

  .page-title {
    font-size: 13px;
    font-weight: 700;
    color: var(--blue);
    letter-spacing: 0.14em;
    text-transform: uppercase;
    margin-bottom: 8px;
  }

  .page-heading {
    font-size: 36px;
    font-weight: 800;
    color: var(--navy);
    line-height: 1.15;
    margin-bottom: 8px;
    letter-spacing: -0.01em;
  }

  .page-sub {
    font-size: 13px;
    color: var(--muted);
    font-weight: 500;
    margin-bottom: 36px;
    letter-spacing: 0.02em;
  }

  .upload-card {
    background: var(--navy);
    border-radius: 12px;
    padding: 28px;
    box-shadow: 0 8px 40px rgba(26,47,94,0.25);
    max-width: 460px;
  }

  .card-title {
    font-size: 13px;
    font-weight: 700;
    color: var(--cyan);
    letter-spacing: 0.12em;
    text-transform: uppercase;
    margin-bottom: 16px;
    display: flex;
    align-items: center;
    gap: 8px;
  }

  .card-title-icon {
    width: 28px; height: 28px;
    background: rgba(77,200,232,0.15);
    border-radius: 6px;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 16px;
  }

  .drop-zone {
    border: 2px dashed rgba(77,200,232,0.4);
    border-radius: 8px;
    padding: 40px 24px;
    text-align: center;
    cursor: pointer;
    transition: all 0.2s;
    background: rgba(255,255,255,0.04);
    position: relative;
    overflow: hidden;
    margin-bottom: 16px;
  }

  .drop-zone.dragover { border-color: var(--cyan); background: rgba(77,200,232,0.08); }
  .drop-zone.has-file { border-color: var(--cyan); border-style: solid; background: rgba(77,200,232,0.06); }

  .drop-icon { font-size: 32px; margin-bottom: 12px; display: block; line-height: 1; }
  .drop-title { font-size: 15px; font-weight: 700; color: var(--white); margin-bottom: 4px; }
  .drop-sub { font-size: 11px; color: var(--muted); font-weight: 500; letter-spacing: 0.04em; }

  .file-types { display: flex; justify-content: center; gap: 6px; margin-top: 12px; }

  .type-badge {
    font-size: 10px; font-weight: 700; padding: 2px 8px;
    border-radius: 3px; letter-spacing: 0.08em; text-transform: uppercase;
  }
  .type-pdf  { background: rgba(192,57,43,0.2); color: #e8a09a; border: 1px solid rgba(192,57,43,0.3); }
  .type-docx { background: rgba(77,200,232,0.15); color: var(--cyan); border: 1px solid rgba(77,200,232,0.3); }
  .type-xlsx { background: rgba(45,122,79,0.2); color: #7fd4a0; border: 1px solid rgba(45,122,79,0.3); }

  .drop-zone input[type="file"] {
    position: absolute; inset: 0; opacity: 0; cursor: pointer; width: 100%; height: 100%;
  }

  .file-preview {
    display: none;
    background: rgba(255,255,255,0.06);
    border: 1px solid rgba(77,200,232,0.25);
    border-radius: 6px;
    padding: 10px 14px;
    align-items: center;
    justify-content: space-between;
    gap: 10px;
    margin-bottom: 14px;
  }
  .file-preview.visible { display: flex; }
  .file-info { display: flex; align-items: center; gap: 10px; }
  .file-icon-wrap { font-size: 18px; }
  .file-name { font-size: 12px; color: var(--white); font-weight: 500; }

  .file-clear {
    background: none; border: 1px solid rgba(255,255,255,0.15); border-radius: 4px;
    color: var(--muted); cursor: pointer; font-size: 11px; font-family: var(--font);
    padding: 3px 8px; transition: all 0.2s; font-weight: 600;
  }
  .file-clear:hover { border-color: var(--danger); color: #e8a09a; }

  .submit-wrap { display: none; }
  .submit-wrap.visible { display: block; }

  .btn-submit {
    width: 100%; background: var(--blue); border: none; border-radius: 6px;
    color: var(--white); cursor: pointer; font-family: var(--font); font-size: 13px;
    font-weight: 700; letter-spacing: 0.1em; padding: 14px 24px;
    transition: all 0.2s; text-transform: uppercase;
  }
  .btn-submit:hover { background: var(--lblue); }
  .btn-submit:disabled { opacity: 0.4; cursor: not-allowed; }

  .status {
    display: none; margin-top: 14px; background: rgba(255,255,255,0.06);
    border: 1px solid rgba(77,200,232,0.2); border-radius: 6px;
    padding: 14px 18px; align-items: center; gap: 12px;
  }
  .status.visible { display: flex; }

  .spinner {
    width: 18px; height: 18px; border: 2px solid rgba(255,255,255,0.15);
    border-top-color: var(--cyan); border-radius: 50%;
    animation: spin 0.8s linear infinite; flex-shrink: 0;
  }
  @keyframes spin { to { transform: rotate(360deg); } }
  .status-text { font-size: 12px; color: var(--muted); font-weight: 500; }

  .error-box {
    display: none; margin-top: 14px; background: rgba(192,57,43,0.12);
    border: 1px solid rgba(192,57,43,0.3); border-radius: 6px;
    padding: 12px 16px; font-size: 12px; color: #e8a09a; font-weight: 500;
  }

  .disclaimer {
    display: none; margin-top: 14px; background: rgba(255,255,255,0.06);
    border: 1px solid rgba(77,200,232,0.15); border-left: 3px solid var(--cyan);
    border-radius: 4px; padding: 10px 14px; font-size: 10px; color: var(--muted);
    line-height: 1.7; font-weight: 500;
  }

  .db-banner {
    display: none; margin-top: 14px; border-radius: 6px;
    padding: 12px 16px; font-size: 12px; font-weight: 600;
    align-items: center; gap: 10px;
  }
  .db-banner.visible { display: flex; }
  .db-success { background: rgba(45,122,79,0.15); border: 1px solid rgba(45,122,79,0.3); color: #7fd4a0; }
  .db-error   { background: rgba(192,57,43,0.12); border: 1px solid rgba(192,57,43,0.3); color: #e8a09a; }
  .db-warning { background: rgba(183,119,13,0.15); border: 1px solid rgba(183,119,13,0.3); color: #f0c060; }

  #results { margin-top: 40px; display: none; max-width: 1100px; }

  .results-header {
    display: flex; align-items: center; justify-content: space-between; margin-bottom: 20px;
  }

  .results-heading { font-size: 22px; font-weight: 800; color: var(--navy); }

  .btn-group { display: flex; gap: 10px; }

  .download-btn {
    display: inline-flex; align-items: center; gap: 8px; background: var(--navy);
    border: none; border-radius: 6px; color: var(--white); cursor: pointer;
    font-family: var(--font); font-size: 11px; font-weight: 700; padding: 10px 20px;
    transition: all 0.2s; text-decoration: none; letter-spacing: 0.08em; text-transform: uppercase;
  }
  .download-btn:hover { background: var(--blue); }
  .download-btn.csv { background: var(--success); }
  .download-btn.csv:hover { background: #3a9e68; }

  .manual-review-banner {
    display: none; margin-bottom: 20px; background: #fffbf0;
    border: 1px solid #e8c84a; border-left: 4px solid #e8c84a;
    border-radius: 6px; padding: 14px 18px; font-size: 12px; color: #7a5c00;
    font-weight: 500; line-height: 1.6;
  }
  .manual-review-banner.visible { display: block; }

  .stats-row { display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin-bottom: 24px; }

  .stat-card {
    background: var(--white); border: 1px solid var(--border); border-radius: 8px;
    padding: 18px 20px; text-align: center; box-shadow: 0 2px 8px rgba(26,47,94,0.06);
  }
  .stat-card.warning-card { border-color: #e8c84a; background: #fffdf0; }

  .stat-number { font-size: 34px; font-weight: 900; color: var(--navy); line-height: 1; }
  .stat-number.warning { color: var(--warning); }
  .stat-label { font-size: 10px; color: var(--muted); margin-top: 6px; text-transform: uppercase; letter-spacing: 0.1em; font-weight: 600; }

  .category-section {
    background: var(--white); border: 1px solid var(--border); border-radius: 8px;
    margin-bottom: 10px; overflow: hidden; box-shadow: 0 2px 8px rgba(26,47,94,0.05);
  }

  .category-header {
    padding: 14px 20px; font-size: 11px; font-weight: 700;
    display: flex; align-items: center; gap: 10px; cursor: pointer;
    user-select: none; border-bottom: 1px solid var(--border);
    background: #f0f6fb; color: var(--navy); letter-spacing: 0.08em;
    text-transform: uppercase; transition: background 0.15s;
  }
  .category-header:hover { background: #e2eef8; }

  .cat-count {
    font-size: 10px; background: var(--navy); color: var(--white);
    border-radius: 3px; padding: 2px 8px; font-weight: 700;
  }

  .cat-toggle { margin-left: auto; color: var(--muted); transition: transform 0.2s; }
  .category-section.collapsed .cat-toggle { transform: rotate(-90deg); }
  .category-section.collapsed .cat-body { display: none; }

  table { width: 100%; border-collapse: collapse; }

  th {
    text-align: left; padding: 10px 16px; font-size: 10px; text-transform: uppercase;
    letter-spacing: 0.1em; color: var(--muted); background: #f7fafd;
    border-bottom: 1px solid var(--border); font-weight: 700;
  }

  td {
    padding: 12px 16px; font-size: 12px; border-bottom: 1px solid #edf3f8;
    vertical-align: top; line-height: 1.5; color: var(--text);
  }

  tr:last-child td { border-bottom: none; }
  tr:hover td { background: #f7fafd; }
  tr.needs-review td { background: #fffdf0; }
  tr.needs-review:hover td { background: #fff8e0; }

  .std-id { font-weight: 700; color: var(--navy); font-size: 12px; }

  .date-chip {
    font-size: 11px; background: #e8f1fa; color: var(--blue);
    padding: 3px 8px; border-radius: 3px; white-space: nowrap;
    font-weight: 600; border: 1px solid var(--border);
  }
  .date-chip.manual { background: #fffbf0; color: var(--warning); border-color: #e8c84a; }

  .badge {
    font-size: 10px; font-weight: 700; padding: 3px 8px;
    border-radius: 3px; letter-spacing: 0.06em; text-transform: uppercase;
  }
  .badge-us   { background: rgba(26,47,94,0.08); color: var(--navy); border: 1px solid rgba(26,47,94,0.15); }
  .badge-intl { background: rgba(45,122,79,0.08); color: var(--success); border: 1px solid rgba(45,122,79,0.2); }

  .review-flag {
    font-size: 10px; font-weight: 700; color: var(--warning);
    background: #fffbf0; border: 1px solid #e8c84a;
    padding: 2px 7px; border-radius: 3px; white-space: nowrap;
  }

  footer {
    position: relative; z-index: 5; background: var(--navy);
    padding: 14px 56px; text-align: center; font-size: 10px;
    color: rgba(255,255,255,0.4); letter-spacing: 0.08em;
    text-transform: uppercase; font-weight: 500; margin-top: auto;
  }
</style>
</head>
<body>
<div class="page">

  <svg class="wave-bg" viewBox="0 0 600 900" xmlns="http://www.w3.org/2000/svg" preserveAspectRatio="xMidYMid slice">
    <defs>
      <linearGradient id="wv1" x1="0%" y1="0%" x2="100%" y2="100%">
        <stop offset="0%" style="stop-color:#5ba3d9;stop-opacity:0.5"/>
        <stop offset="100%" style="stop-color:#2e6db4;stop-opacity:0.3"/>
      </linearGradient>
      <linearGradient id="wv2" x1="0%" y1="0%" x2="100%" y2="100%">
        <stop offset="0%" style="stop-color:#4dc8e8;stop-opacity:0.4"/>
        <stop offset="100%" style="stop-color:#5ba3d9;stop-opacity:0.2"/>
      </linearGradient>
      <linearGradient id="wv3" x1="0%" y1="0%" x2="100%" y2="100%">
        <stop offset="0%" style="stop-color:#7b68ee;stop-opacity:0.3"/>
        <stop offset="100%" style="stop-color:#2e6db4;stop-opacity:0.15"/>
      </linearGradient>
    </defs>
    <path d="M600,0 L600,900 L0,900 Q150,700 300,600 Q450,500 600,300 Z" fill="url(#wv3)" opacity="0.5"/>
    <path d="M600,0 L600,900 L100,900 Q200,750 350,650 Q500,550 600,350 Z" fill="url(#wv1)" opacity="0.6"/>
    <path d="M600,200 Q500,350 400,450 Q300,550 250,700 Q200,800 300,900 L600,900 Z" fill="url(#wv2)" opacity="0.7"/>
    <ellipse cx="500" cy="200" rx="120" ry="80" fill="#4dc8e8" opacity="0.15" transform="rotate(-20 500 200)"/>
  </svg>

  <header>
    <div class="logo-wrap">
      <img class="logo-img"
        src="https://static.wixstatic.com/media/9f2dc8_980147f3f25a4e50a2220ab0bd98dba8~mv2.png/v1/fill/w_129,h_128,al_c,q_85,usm_0.66_1.00_0.01,enc_avif,quality_auto/Sig%20logo.png"
        alt="Rallis-Daw Consulting">
    </div>
    <div class="tool-label">RegCheck &mdash; Prototype v0.1</div>
  </header>

  <div class="content">

    <div class="page-title">Regulatory Standards</div>
    <div class="page-heading">External Documents<br>Checker</div>
    <div class="page-sub">Upload a device requirements file to extract and review all regulatory standards</div>

    <div class="upload-card">

      <div class="card-title">
        <div class="card-title-icon">&#128203;</div>
        Document Upload
      </div>

      <div class="drop-zone" id="dropZone">
        <input type="file" id="fileInput" accept=".pdf,.docx,.xlsx">
        <span class="drop-icon" id="dropIcon">&#9729;</span>
        <div class="drop-title" id="dropTitle">Drag &amp; Drop or Upload a File</div>
        <div class="drop-sub" id="dropSub">Click anywhere in this box to browse</div>
        <div class="file-types">
          <span class="type-badge type-pdf">PDF</span>
          <span class="type-badge type-docx">DOCX</span>
          <span class="type-badge type-xlsx">XLSX</span>
        </div>
      </div>

      <div class="file-preview" id="filePreview">
        <div class="file-info">
          <span class="file-icon-wrap" id="fileTypeIcon">&#128196;</span>
          <span class="file-name" id="fileName"></span>
        </div>
        <button class="file-clear" onclick="clearFile()">Remove</button>
      </div>

      <div class="submit-wrap" id="submitWrap">
        <button class="btn-submit" id="submitBtn" onclick="submitFile()">
          Extract Requirements
        </button>
      </div>

      <div class="status" id="status">
        <div class="spinner"></div>
        <div class="status-text" id="statusText">Sending to Gemini AI &mdash; this may take 15&ndash;30 seconds...</div>
      </div>

      <div class="error-box" id="errorBox"></div>

      <div class="db-banner db-success" id="dbSuccess">
        &#10003; <span id="dbSuccessMsg">Standards saved to database successfully.</span>
      </div>
      <div class="db-banner db-error" id="dbError">
        &#9888; <span id="dbErrorMsg">Could not save to database.</span>
      </div>
      <div class="db-banner db-warning" id="dbWarning">
        &#9888; Supabase is not configured. Results will not be saved to the database.
      </div>

      <div class="disclaimer" id="disclaimer">
        AI DISCLAIMER &mdash; This extraction is intended solely to assist regulatory review.
        All results must be verified by a qualified professional before use in any compliance submission.
      </div>

    </div>

    <div id="results">
      <div class="results-header">
        <div class="results-heading">Extraction Results</div>
        <div class="btn-group">
          <a class="download-btn csv" id="downloadCsv" href="#">Download CSV</a>
          <a class="download-btn" id="downloadJson" href="#">Download JSON</a>
        </div>
      </div>

      <div class="manual-review-banner" id="manualReviewBanner">
        <strong>Manual Review Required</strong> &mdash; Some standards are marked as
        <strong>"Current"</strong> because the original document used <strong>**</strong>
        instead of a specific date. These standards cannot be automatically compared and
        must be verified manually by visiting the relevant regulatory body's website.
      </div>

      <div class="stats-row" id="statsRow"></div>
      <div id="categorySections"></div>
    </div>

  </div>

  <footer>
    &copy; Rallis-Daw Consulting LLC &nbsp;&middot;&nbsp; RegCheck Tool &nbsp;&middot;&nbsp; For internal use only
  </footer>

</div>
<script>
  var dropZone    = document.getElementById('dropZone');
  var fileInput   = document.getElementById('fileInput');
  var filePreview = document.getElementById('filePreview');
  var fileName    = document.getElementById('fileName');
  var submitWrap  = document.getElementById('submitWrap');
  var submitBtn   = document.getElementById('submitBtn');
  var statusBox   = document.getElementById('status');
  var results     = document.getElementById('results');
  var errorBox    = document.getElementById('errorBox');
  var disclaimer  = document.getElementById('disclaimer');
  var dbSuccess   = document.getElementById('dbSuccess');
  var dbError     = document.getElementById('dbError');
  var dbWarning   = document.getElementById('dbWarning');

  var selectedFile = null;

  var ALLOWED = ['pdf', 'docx', 'xlsx'];
  var FILE_ICONS = { pdf: '&#128212;', docx: '&#128216;', xlsx: '&#128218;' };

  function getExt(filename) {
    return filename.split('.').pop().toLowerCase();
  }

  dropZone.addEventListener('dragover', function(e) {
    e.preventDefault();
    dropZone.classList.add('dragover');
  });

  dropZone.addEventListener('dragleave', function() {
    dropZone.classList.remove('dragover');
  });

  dropZone.addEventListener('drop', function(e) {
    e.preventDefault();
    dropZone.classList.remove('dragover');
    var file = e.dataTransfer.files[0];
    if (file && ALLOWED.indexOf(getExt(file.name)) !== -1) {
      setFile(file);
    } else {
      showError('Please drop a PDF, DOCX, or XLSX file.');
    }
  });

  fileInput.addEventListener('change', function() {
    if (fileInput.files[0]) {
      setFile(fileInput.files[0]);
    }
  });

  function setFile(file) {
    selectedFile = file;
    var ext = getExt(file.name);
    dropZone.classList.add('has-file');
    document.getElementById('dropIcon').textContent = '';
    document.getElementById('dropIcon').innerHTML = '&#10003;';
    document.getElementById('dropTitle').textContent = 'File ready';
    document.getElementById('dropSub').textContent = 'Drop a different file to replace it';
    document.getElementById('fileTypeIcon').innerHTML = FILE_ICONS[ext] || '&#128196;';
    fileName.textContent = file.name;
    filePreview.classList.add('visible');
    submitWrap.classList.add('visible');
    errorBox.style.display = 'none';
    results.style.display = 'none';
    disclaimer.style.display = 'none';
    hideDbBanners();
  }

  function clearFile() {
    selectedFile = null;
    fileInput.value = '';
    dropZone.classList.remove('has-file');
    document.getElementById('dropIcon').innerHTML = '&#9729;';
    document.getElementById('dropTitle').textContent = 'Drag & Drop or Upload a File';
    document.getElementById('dropSub').textContent = 'Click anywhere in this box to browse';
    filePreview.classList.remove('visible');
    submitWrap.classList.remove('visible');
    errorBox.style.display = 'none';
    hideDbBanners();
  }

  function hideDbBanners() {
    dbSuccess.classList.remove('visible');
    dbError.classList.remove('visible');
    dbWarning.classList.remove('visible');
  }

  function submitFile() {
    if (!selectedFile) return;
    submitBtn.disabled = true;
    submitBtn.textContent = 'Processing...';
    errorBox.style.display = 'none';
    results.style.display = 'none';
    disclaimer.style.display = 'none';
    hideDbBanners();
    statusBox.classList.add('visible');

    var formData = new FormData();
    formData.append('file', selectedFile);

    fetch('/extract', { method: 'POST', body: formData })
      .then(function(r) { return r.json(); })
      .then(function(data) {
        statusBox.classList.remove('visible');
        submitBtn.disabled = false;
        submitBtn.textContent = 'Extract Requirements';

        if (data.error) {
          showError(data.error);
          return;
        }

        if (data.db_success) {
          document.getElementById('dbSuccessMsg').textContent =
            data.db_count + ' standards saved to database successfully.';
          dbSuccess.classList.add('visible');
        } else if (data.db_error) {
          document.getElementById('dbErrorMsg').textContent =
            'Could not save to database: ' + data.db_error;
          dbError.classList.add('visible');
        } else if (data.db_not_configured) {
          dbWarning.classList.add('visible');
        }

        disclaimer.style.display = 'block';
        renderResults(data.requirements, data.filename);
      })
      .catch(function(err) {
        statusBox.classList.remove('visible');
        submitBtn.disabled = false;
        submitBtn.textContent = 'Extract Requirements';
        showError('Server error: ' + err.message);
      });
  }

  function renderResults(requirements, filename) {
    results.style.display = 'block';

    var manualCount = requirements.filter(function(r) { return r.needs_manual_review; }).length;

    var banner = document.getElementById('manualReviewBanner');
    if (manualCount > 0) {
      banner.classList.add('visible');
    } else {
      banner.classList.remove('visible');
    }

    var categories = {};
    requirements.forEach(function(r) {
      var cat = r.category || 'Uncategorized';
      if (!categories[cat]) categories[cat] = [];
      categories[cat].push(r);
    });

    var usCount   = requirements.filter(function(r) { return r.region === 'US'; }).length;
    var intlCount = requirements.filter(function(r) { return r.region === 'International'; }).length;

    var warnClass  = manualCount > 0 ? ' warning-card' : '';
    var warnNum    = manualCount > 0 ? ' warning' : '';

    document.getElementById('statsRow').innerHTML =
      '<div class="stat-card"><div class="stat-number">' + requirements.length + '</div><div class="stat-label">Total Standards</div></div>' +
      '<div class="stat-card"><div class="stat-number">' + Object.keys(categories).length + '</div><div class="stat-label">Categories</div></div>' +
      '<div class="stat-card"><div class="stat-number">' + usCount + '</div><div class="stat-label">US Standards</div></div>' +
      '<div class="stat-card' + warnClass + '"><div class="stat-number' + warnNum + '">' + manualCount + '</div><div class="stat-label">Verify Manually</div></div>';

    var jsonBlob = new Blob([JSON.stringify(requirements, null, 2)], { type: 'application/json' });
    var jsonBtn  = document.getElementById('downloadJson');
    jsonBtn.href = URL.createObjectURL(jsonBlob);
    jsonBtn.download = filename.replace(/[.](pdf|docx|xlsx)$/i, '_requirements.json');

    var csvBtn = document.getElementById('downloadCsv');
    csvBtn.onclick = function(e) {
      e.preventDefault();
      var rows = [['ID','Standard ID','Date','Date Year','Category','Region','Description','Needs Manual Review']];
      requirements.forEach(function(r) {
        rows.push([
          r.id,
          r.standard_id,
          r.date,
          r.date_year || '',
          r.category,
          r.region,
          r.description,
          r.needs_manual_review ? 'Yes' : 'No'
        ]);
      });
      var csv = rows.map(function(row) {
        return row.map(function(v) {
          return '"' + String(v).replace(/"/g, '""') + '"';
        }).join(',');
      }).join('\n');
      var blob = new Blob([csv], { type: 'text/csv' });
      var url  = URL.createObjectURL(blob);
      var a    = document.createElement('a');
      a.href   = url;
      a.download = filename.replace(/[.](pdf|docx|xlsx)$/i, '_requirements.csv');
      a.click();
    };

    var container = document.getElementById('categorySections');
    container.innerHTML = '';

    Object.keys(categories).forEach(function(cat) {
      var items = categories[cat];
      var rows  = items.map(function(item) {
        var isManual    = item.needs_manual_review;
        var dateClass   = isManual ? 'date-chip manual' : 'date-chip';
        var dateDisplay = isManual ? item.date + ' *' : item.date;
        var reviewFlag  = isManual ? '<span class="review-flag">Verify Manually</span>' : '';
        var regionClass = item.region === 'US' ? 'badge-us' : 'badge-intl';
        return '<tr class="' + (isManual ? 'needs-review' : '') + '">' +
          '<td><span class="std-id">' + (item.standard_id || '') + '</span></td>' +
          '<td><span class="' + dateClass + '">' + dateDisplay + '</span></td>' +
          '<td><span class="badge ' + regionClass + '">' + (item.region || '') + '</span></td>' +
          '<td style="font-size:12px;color:#4a5568">' + (item.description || '') + '</td>' +
          '<td>' + reviewFlag + '</td>' +
          '</tr>';
      }).join('');

      var section = document.createElement('div');
      section.className = 'category-section';
      section.innerHTML =
        '<div class="category-header" onclick="this.parentElement.classList.toggle(\'collapsed\')">' +
          '<span>' + cat + '</span>' +
          '<span class="cat-count">' + items.length + '</span>' +
          '<span class="cat-toggle">&#9660;</span>' +
        '</div>' +
        '<div class="cat-body">' +
          '<table>' +
            '<thead><tr><th>Standard ID</th><th>Date</th><th>Region</th><th>Description</th><th>Status</th></tr></thead>' +
            '<tbody>' + rows + '</tbody>' +
          '</table>' +
        '</div>';
      container.appendChild(section);
    });

    results.scrollIntoView({ behavior: 'smooth', block: 'start' });
  }

  function showError(msg) {
    errorBox.textContent = 'Error: ' + msg;
    errorBox.style.display = 'block';
  }
</script>
</body>
</html>"""


@app.route("/")
def index():
    return render_template_string(HTML)


@app.route("/extract", methods=["POST"])
def extract():
    if not GEMINI_API_KEY:
        return jsonify({"error": "No Gemini API key found. Check your .env file."})

    uploaded = request.files.get("file")
    if not uploaded:
        return jsonify({"error": "No file received."})

    filename = uploaded.filename
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext not in ("pdf", "docx", "xlsx"):
        return jsonify({"error": "Unsupported file type: ." + ext + ". Please upload a PDF, DOCX, or XLSX."})

    with tempfile.NamedTemporaryFile(suffix="." + ext, delete=False) as tmp:
        uploaded.save(tmp.name)
        tmp_path = tmp.name

    try:
        text = extract_text(tmp_path, filename)
        if not text.strip():
            return jsonify({"error": "Could not extract text from the file."})

        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(GEMINI_MODEL)

        prompt = (
            "You are a regulatory document parser for medical devices.\n\n"
            "Extract every standard, regulation, or guidance document from the text below.\n"
            "Return ONLY a valid JSON array with no preamble, explanation, or markdown fences.\n\n"
            "Each object must have exactly these fields:\n"
            "- standard_id: the identifier e.g. ISO 13485, 21 CFR Part 820, FDA Guidance 1677\n"
            "- date: the revision or issue date as written. If marked with ** use the word Current\n"
            "- category: the section this belongs to e.g. Quality System Management\n"
            "- region: US or International based on context\n"
            "- description: a brief description of what this standard covers\n\n"
            "Document text:\n" + text
        )

        response = model.generate_content(prompt)
        raw = response.text.strip()

        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        requirements = json.loads(raw)
        enriched = [enrich_requirement(req, filename) for req in requirements]

        for i, req in enumerate(enriched):
            req["id"] = i + 1

        db_success        = False
        db_error          = None
        db_count          = 0
        db_not_configured = False

        if not supabase_client:
            db_not_configured = True
        else:
            rows_to_insert = [{k: v for k, v in r.items() if k != "id"} for r in enriched]
            db_count, db_error = insert_to_supabase(rows_to_insert)
            db_success = db_error is None

        return jsonify({
            "requirements":      enriched,
            "filename":          filename,
            "db_success":        db_success,
            "db_count":          db_count,
            "db_error":          db_error,
            "db_not_configured": db_not_configured,
        })

    except json.JSONDecodeError:
        return jsonify({"error": "Gemini returned an invalid response. Try running again."})
    except Exception as e:
        return jsonify({"error": str(e)})
    finally:
        os.unlink(tmp_path)


if __name__ == "__main__":
    print("\n" + "=" * 50)
    print("  RegCheck -- Rallis-Daw Consulting")
    print("  Supports: PDF, DOCX, XLSX")
    print("=" * 50)

    if not GEMINI_API_KEY:
        print("\n  WARNING: No Gemini API key found!")
        print("  Add GEMINI_API_KEY to your .env file\n")
    else:
        print("  OK: Gemini API key loaded")

    if not SUPABASE_URL or not SUPABASE_KEY:
        print("  WARNING: Supabase not configured!")
        print("  Add SUPABASE_URL and SUPABASE_KEY to your .env file")
    else:
        print("  OK: Supabase configured")

    print("\n  Open your browser to: http://localhost:5000")
    print("  Press Ctrl+C to stop the server")
    print("=" * 50 + "\n")

    app.run(debug=False, port=5000)
