"""
Medical Device Requirements Extractor - Local Web App
Rallis-Daw Consulting | RegCheck Tool
------------------------------------------------------
Supports PDF, Word (.docx), and Excel (.xlsx) files.

SETUP:
    pip3 install flask google-generativeai pdfplumber python-docx openpyxl python-dotenv

CREATE A .env FILE in this same folder containing:
    GEMINI_API_KEY=your_api_key_here

RUN:
    python3 app.py

Then open your browser to: http://localhost:5000
"""

import json
import os
import tempfile
import pdfplumber
import google.generativeai as genai
from flask import Flask, request, jsonify, render_template
from docx import Document
from dotenv import load_dotenv
import openpyxl

# New imports for Firebase Admin SDK and Supabase
import firebase_admin
from firebase_admin import credentials, auth as firebase_auth
from flask_cors import CORS
from supabase import create_client
from functools import wraps

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL   = "gemini-2.5-flash"

# ── NEW: Firebase init ─────────────────────────────────────────────────────────
cert = json.loads(os.environ.get("FIREBASE_CREDENTIALS"))
cred = credentials.Certificate(cert)
firebase_admin.initialize_app(cred)

# ── NEW: Supabase init ─────────────────────────────────────────────────────────
supabase = create_client(
  os.environ.get("SUPABASE_URL"),
  os.environ.get("SUPABASE_KEY")
)

app = Flask(__name__, template_folder='templates', static_folder='static')

CORS(app, origins=["http://localhost:5000"])  # UPDATE after deployment

# ── TEXT EXTRACTION ────────────────────────────────────────────────────────────

def extract_from_pdf(path):
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_from_docx(path):
    doc = Document(path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += row_text + "\n"
    return text


def extract_from_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    text = ""
    for sheet in wb.worksheets:
        text += f"\n--- Sheet: {sheet.title} ---\n"
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                text += row_text + "\n"
    return text


def extract_text(path, filename):
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_from_pdf(path)
    elif ext == "docx":
        return extract_from_docx(path)
    elif ext == "xlsx":
        return extract_from_xlsx(path)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/extract", methods=["POST"])
def extract():
    if not GEMINI_API_KEY:
        return jsonify({"error": "No API key found. Make sure your .env file exists with GEMINI_API_KEY set."})

    uploaded = request.files.get("file")
    if not uploaded:
        return jsonify({"error": "No file received."})

    filename = uploaded.filename
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext not in ("pdf", "docx", "xlsx"):
        return jsonify({"error": f"Unsupported file type: .{ext}. Please upload a PDF, DOCX, or XLSX file."})

    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
        uploaded.save(tmp.name)
        tmp_path = tmp.name

    try:
        text = extract_text(tmp_path, filename)

        if not text.strip():
            return jsonify({"error": "Could not extract text from the file."})

        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(GEMINI_MODEL)

        prompt = f"""
You are a regulatory document parser for medical devices.

Extract every standard, regulation, or guidance document from the text below.
Return ONLY a valid JSON array with no preamble, explanation, or markdown fences.

Each object must have exactly these fields:
- "standard_id": the identifier (e.g. "ISO 13485", "21 CFR Part 820", "FDA Guidance 1677")
- "date": the revision or issue date as written. If marked with ** it means check website for current version — use "Current"
- "category": the section this belongs to (e.g. "Quality System Management", "Biocompatibility")
- "region": "US" or "International" based on context
- "description": a brief description of what this standard covers

Document text:
{text}
"""

        response = model.generate_content(prompt)
        raw = response.text.strip()

        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        requirements = json.loads(raw)

        for i, req in enumerate(requirements):
            req["id"] = i + 1

        return jsonify({
            "requirements": requirements,
            "filename": filename
        })

    except json.JSONDecodeError:
        return jsonify({"error": "Gemini returned an invalid response. Try running again."})
    except Exception as e:
        return jsonify({"error": str(e)})
    finally:
        os.unlink(tmp_path)


if __name__ == "__main__":
    print("\n" + "=" * 50)
    print("  RegCheck — Rallis-Daw Consulting")
    print("  Supports: PDF, DOCX, XLSX")
    print("=" * 50)

    if not GEMINI_API_KEY:
        print("\n  ⚠ WARNING: No API key found!")
        print("  Create a .env file with: GEMINI_API_KEY=your_key_here\n")
    else:
        print("  ✓ API key loaded from .env")

    print("\n  Open your browser to: http://localhost:5000")
    print("  Press Ctrl+C to stop the server")
    print("=" * 50 + "\n")

    app.run(debug=False, port=5000)