import pdfplumber
from docx import Document
import openpyxl

def extract_from_pdf(path):
    """Extract text from PDF files."""
    text = ""

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    return text

def extract_from_docx(path):
    """Extract text from Word documents."""
    doc = Document(path)
    text = ""

    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += row_text + "\n"

    return text

def extract_from_xlsx(path):
    """Extract text from Excel files."""
    wb = openpyxl.load_workbook(path, data_only=True)
    text = ""

    for sheet in wb.worksheets:
        text += f"\n--- Sheet: {sheet.title} ---\n"

        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) for cell in row if cell is not None)

            if row_text.strip():
                text += row_text + "\n"

    return text

def extract_text(path, filename):
    """Main extraction function - routes to appropriate extractor."""
    ext = filename.rsplit(".", 1)[-1].lower()
    
    if ext == "pdf":
        return extract_from_pdf(path)
    elif ext == "docx":
        return extract_from_docx(path)
    elif ext == "xlsx":
        return extract_from_xlsx(path)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")